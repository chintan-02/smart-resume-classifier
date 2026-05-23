from pathlib import Path
from io import BytesIO
import re
from typing import Dict, List, Optional

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover - handled gracefully at runtime.
    PdfReader = None

try:
    from docx import Document
except ImportError:  # pragma: no cover - handled gracefully at runtime.
    Document = None

from src.app_config import SUPPORTED_FILE_EXTENSIONS


EMPTY_SECTIONS = {
    "summary": "",
    "skills": "",
    "experience": "",
    "projects": "",
    "education": "",
    "certifications": "",
    "volunteering": "",
    "achievements": "",
}

SECTION_ALIASES = {
    "summary": "summary",
    "professional summary": "summary",
    "career profile": "summary",
    "profile": "summary",
    "objective": "summary",
    "technical skills": "skills",
    "skills": "skills",
    "core skills": "skills",
    "professional skills": "skills",
    "professional experience": "experience",
    "work experience": "experience",
    "experience": "experience",
    "employment history": "experience",
    "projects": "projects",
    "project experience": "projects",
    "education": "education",
    "academic background": "education",
    "certifications": "certifications",
    "certification": "certifications",
    "professional development": "certifications",
    "training": "certifications",
    "volunteering": "volunteering",
    "volunteer experience": "volunteering",
    "achievements": "achievements",
    "accomplishments": "achievements",
}

TEMPLATE_PLACEHOLDERS = [
    "your name",
    "professional title",
    "job title",
    "company name",
    "year to year",
    "month, year",
    "email address",
    "phone number",
    "linkedin url",
    "github/portfolio url",
    "result oriented statement",
    "the name of the program",
    "university name",
    "institution name",
    "graduation year",
    "add measurable result",
    "placeholder",
]

STRONG_TEMPLATE_WARNING = (
    "This appears to be a resume template or incomplete resume. "
    "Please upload a completed resume for accurate scoring."
)
PARTIAL_TEMPLATE_WARNING = (
    "This resume contains some placeholder fields. The analysis can continue, "
    "but replacing placeholders with real details will improve scoring accuracy."
)


def _reset_file(uploaded_file) -> None:
    if hasattr(uploaded_file, "seek"):
        uploaded_file.seek(0)


def _read_uploaded_bytes(uploaded_file) -> bytes:
    _reset_file(uploaded_file)
    if hasattr(uploaded_file, "getvalue"):
        return uploaded_file.getvalue()
    data = uploaded_file.read()
    if isinstance(data, str):
        return data.encode("utf-8", errors="ignore")
    return data or b""


def extract_text_from_pdf(uploaded_file) -> str:
    try:
        if PdfReader is None:
            return ""
        _reset_file(uploaded_file)
        reader = PdfReader(uploaded_file)
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    except Exception:
        return ""


def extract_text_from_txt(uploaded_file) -> str:
    try:
        return _read_uploaded_bytes(uploaded_file).decode("utf-8", errors="ignore")
    except Exception:
        return ""


def extract_text_from_docx(uploaded_file) -> str:
    try:
        if Document is None:
            return ""

        docx_bytes = _read_uploaded_bytes(uploaded_file)
        if not docx_bytes:
            return ""

        document = Document(BytesIO(docx_bytes))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        table_cells = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        table_cells.append(cell_text)
        return "\n".join(paragraphs + table_cells).strip()
    except Exception:
        return ""


def get_file_extension(uploaded_file) -> str:
    return Path(getattr(uploaded_file, "name", "")).suffix.lower()


def is_supported_file(uploaded_file) -> bool:
    return get_file_extension(uploaded_file) in SUPPORTED_FILE_EXTENSIONS


def extract_resume_text(uploaded_file) -> str:
    file_name = getattr(uploaded_file, "name", "").lower()
    if file_name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    if file_name.endswith(".txt"):
        return extract_text_from_txt(uploaded_file)
    if file_name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    return ""


def _normalize_heading(text: str) -> str:
    return re.sub(r"[^a-z0-9\s/&-]", "", text.lower()).strip()


def _heading_key(line: str) -> Optional[str]:
    cleaned = _normalize_heading(line)
    cleaned = re.sub(r"\s+", " ", cleaned)
    if cleaned in SECTION_ALIASES:
        return SECTION_ALIASES[cleaned]

    if len(cleaned.split()) <= 4:
        for heading, key in SECTION_ALIASES.items():
            if cleaned == heading or cleaned.startswith(f"{heading} "):
                return key
    return None


def _clean_lines(text: str) -> List[str]:
    return [line.strip(" \t-*•|") for line in text.splitlines() if line.strip(" \t-*•|")]


def detect_resume_sections(resume_text: str) -> Dict[str, str]:
    sections = EMPTY_SECTIONS.copy()
    if not resume_text:
        return sections

    current_key = None
    collected = {key: [] for key in sections}

    for raw_line in resume_text.splitlines():
        line = raw_line.strip()
        if not line:
            if current_key:
                collected[current_key].append("")
            continue

        possible_heading = line.rstrip(":").strip()
        heading_key = _heading_key(possible_heading)
        if heading_key:
            current_key = heading_key
            continue

        if current_key:
            collected[current_key].append(line)

    return {key: "\n".join(value).strip() for key, value in collected.items()}


def extract_contact_info(resume_text: str) -> Dict[str, Optional[str]]:
    text = resume_text or ""
    urls = re.findall(r"https?://[^\s,)>\]]+|www\.[^\s,)>\]]+", text, flags=re.IGNORECASE)
    email_match = re.search(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", text, flags=re.IGNORECASE)
    phone_match = re.search(
        r"(?:(?:\+?1[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)\d{3}[\s.-]?\d{4})",
        text,
    )

    linkedin = next((url for url in urls if "linkedin.com" in url.lower()), None)
    github = next((url for url in urls if "github.com" in url.lower()), None)
    portfolio = next(
        (
            url
            for url in urls
            if "linkedin.com" not in url.lower()
            and "github.com" not in url.lower()
            and not re.search(r"\.(pdf|docx?|txt)$", url, flags=re.IGNORECASE)
        ),
        None,
    )

    return {
        "email": email_match.group(0) if email_match else None,
        "phone": phone_match.group(0) if phone_match else None,
        "linkedin": linkedin,
        "github": github,
        "portfolio": portfolio,
    }


def extract_education_info(sections: Dict[str, str], resume_text: str) -> List[str]:
    education_text = sections.get("education", "") if sections else ""
    lines = _clean_lines(education_text)
    if not lines and resume_text:
        degree_pattern = re.compile(r"\b(bachelor|master|phd|degree|university|college|diploma|b\.s\.|m\.s\.|mba)\b", re.I)
        lines = [line for line in _clean_lines(resume_text) if degree_pattern.search(line)]
    return lines[:20]


def extract_experience_info(sections: Dict[str, str], resume_text: str) -> List[str]:
    experience_text = sections.get("experience", "") if sections else ""
    lines = _clean_lines(experience_text)
    if not lines and resume_text:
        experience_pattern = re.compile(r"\b(experience|engineer|developer|analyst|manager|intern|consultant)\b", re.I)
        lines = [line for line in _clean_lines(resume_text) if experience_pattern.search(line)]
    return lines[:40]


def extract_project_info(sections: Dict[str, str], resume_text: str) -> List[str]:
    project_text = sections.get("projects", "") if sections else ""
    lines = _clean_lines(project_text)
    if not lines and resume_text:
        project_pattern = re.compile(r"\b(project|built|developed|implemented|created)\b", re.I)
        lines = [line for line in _clean_lines(resume_text) if project_pattern.search(line)]
    return lines[:30]


def extract_certifications(sections: Dict[str, str], resume_text: str) -> List[str]:
    certifications_text = sections.get("certifications", "") if sections else ""
    lines = _clean_lines(certifications_text)
    if not lines and resume_text:
        cert_pattern = re.compile(r"\b(certification|certified|certificate|professional development|training)\b", re.I)
        lines = [line for line in _clean_lines(resume_text) if cert_pattern.search(line)]
    return lines[:25]


def estimate_years_of_experience(resume_text: str) -> Optional[int]:
    text = resume_text or ""
    matches = re.findall(r"\b(\d{1,2})\s*\+?\s*(?:years|yrs)\b(?:\s+of\s+experience)?", text, flags=re.IGNORECASE)
    if not matches:
        matches = re.findall(r"\b(\d{1,2})\s*\+?\s+years?\s+of\s+(?:professional\s+)?experience\b", text, flags=re.IGNORECASE)
    if not matches:
        return None
    return max(int(value) for value in matches)


def detect_template_resume(resume_text: str) -> dict:
    text = resume_text or ""
    lower_text = text.lower()
    matched = []
    for placeholder in TEMPLATE_PLACEHOLDERS:
        if placeholder in lower_text:
            matched.append(placeholder)

    bracket_patterns = {
        "[number]": r"\[\s*number\s*\]",
        "[location]": r"\[\s*location\s*\]",
        "your name": r"\[\s*your\s+name\s*\]",
        "job title": r"\[\s*job\s+title\s*\]",
        "company name": r"\[\s*company\s+name\s*\]",
        "month, year": r"\[\s*month\s*,?\s*year\s*\]",
        "university name": r"\[\s*university\s+name\s*\]",
        "placeholder": r"\[[^\]]*placeholder[^\]]*\]",
    }
    for label, pattern in bracket_patterns.items():
        if re.search(pattern, lower_text, flags=re.IGNORECASE):
            matched.append(label)

    matched_placeholders = list(dict.fromkeys(matched))
    placeholder_count = len(matched_placeholders)

    sections = detect_resume_sections(text)
    real_content_signals = []

    if sections.get("skills", "").strip() or re.search(r"^\s*(technical\s+skills|skills)\s*:?\s*$", text, flags=re.I | re.M):
        real_content_signals.append("technical skills section exists")

    skill_pattern = re.compile(
        r"\b(python|sql|pandas|numpy|scikit[-\s]?learn|docker|aws|tableau|power\s*bi|machine\s+learning|nlp)\b",
        flags=re.IGNORECASE,
    )
    found_skills = sorted({match.group(0).strip() for match in skill_pattern.finditer(text)}, key=str.lower)
    if found_skills:
        real_content_signals.append("resume skills found: " + ", ".join(found_skills[:8]))

    action_bullet_pattern = re.compile(
        r"(?im)^\s*(?:[-*•]\s*)?(developed|automated|built|engineered|optimized)\b.+"
    )
    action_bullets = action_bullet_pattern.findall(text)
    if action_bullets:
        real_content_signals.append("experience bullets with action verbs")

    measurable_result_pattern = re.compile(
        r"\b(?:\d+(?:\.\d+)?%|\d+\s*\+|\d+[kKmM]\+?|\d{2,})\b"
    )
    if measurable_result_pattern.search(text):
        real_content_signals.append("measurable results or numeric details")

    if sections.get("projects", "").strip() or re.search(r"^\s*projects?\s*:?\s*$", text, flags=re.I | re.M):
        real_content_signals.append("projects section exists")

    if sections.get("education", "").strip() or re.search(r"^\s*education\s*:?\s*$", text, flags=re.I | re.M):
        real_content_signals.append("education section exists")

    real_content_signal_count = len(real_content_signals)

    template_score = max(0, min(100, (placeholder_count * 12) - (real_content_signal_count * 8)))

    if placeholder_count >= 8 and real_content_signal_count < 3:
        severity = "strong"
    elif placeholder_count >= 3 and real_content_signal_count >= 3:
        severity = "partial"
    elif placeholder_count < 3:
        severity = "none"
    else:
        severity = "partial"

    warning = ""
    if severity == "strong":
        warning = STRONG_TEMPLATE_WARNING
    elif severity == "partial":
        warning = PARTIAL_TEMPLATE_WARNING

    return {
        "is_template": severity == "strong",
        "template_score": template_score,
        "severity": severity,
        "matched_placeholders": matched_placeholders,
        "real_content_signals": real_content_signals,
        "warning": warning,
    }


def parse_resume(resume_file) -> dict:
    extracted_text = extract_resume_text(resume_file)
    sections = detect_resume_sections(extracted_text)
    contact_info = extract_contact_info(extracted_text)

    return {
        "text": extracted_text,
        "sections": sections,
        "contact_info": contact_info,
        "education": extract_education_info(sections, extracted_text),
        "experience": extract_experience_info(sections, extracted_text),
        "projects": extract_project_info(sections, extracted_text),
        "certifications": extract_certifications(sections, extracted_text),
        "estimated_years_experience": estimate_years_of_experience(extracted_text),
        "template_detection": detect_template_resume(extracted_text),
    }
